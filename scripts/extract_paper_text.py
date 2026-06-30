"""Extract plain text from .docx and .pdf files for domain term mining.

Outputs two .txt files alongside the originals for review.
"""
import sys
import os
from pathlib import Path

# Ensure UTF-8 on Windows
os.environ.setdefault("PYTHONIOENCODING", "utf-8")


def extract_docx(path: Path) -> str:
    """Extract all text from a .docx file (paragraphs + tables)."""
    import docx
    doc = docx.Document(str(path))
    parts = []

    # Walk paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Walk tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_pdf(path: Path) -> str:
    """Extract all text from a .pdf file (page by page)."""
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    parts = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            parts.append(text.strip())
    doc.close()
    return "\n\n--- PAGE BREAK ---\n\n".join(parts)


def main():
    downloads = Path(os.environ["USERPROFILE"]) / "Downloads"
    out_dir = Path(os.environ["TEMP"]) / "opencode" / "paper_text"
    out_dir.mkdir(parents=True, exist_ok=True)

    files = [
        (downloads / "AOP review v8.docx", "aop_review"),
        (downloads / "s41592-025-02655-w.pdf", "nature_methods"),
    ]

    for src, label in files:
        if not src.exists():
            print(f"SKIP: {src} not found")
            continue

        print(f"\n{'='*60}")
        print(f"Extracting: {src.name}")
        print(f"{'='*60}")

        if src.suffix == ".docx":
            text = extract_docx(src)
        elif src.suffix == ".pdf":
            text = extract_pdf(src)
        else:
            print(f"SKIP: unsupported format {src.suffix}")
            continue

        out_path = out_dir / f"{label}.txt"
        out_path.write_text(text, encoding="utf-8")

        word_count = len(text.split())
        char_count = len(text)
        print(f"  Chars: {char_count:,}")
        print(f"  Words: {word_count:,}")
        print(f"  Saved: {out_path}")

        # Print first 500 chars as preview
        print(f"\n  Preview (first 500 chars):")
        print(f"  {text[:500]}")
        print(f"  ...")
        print(f"\n  Last 300 chars:")
        print(f"  ...{text[-300:]}")


if __name__ == "__main__":
    main()
